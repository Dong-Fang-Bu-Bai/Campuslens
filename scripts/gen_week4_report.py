"""从第四周周报生成 Word 文档。"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

title = doc.add_heading('第四周工作周报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('基本信息', level=2)
table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate([
    ('项目', '内容'), ('姓名', '庄子杰'),
    ('模块', 'M5 用户反馈纠错与检索记录统计'),
    ('日期', '2026-06-08 ~ 2026-06-14'),
    ('分支', 'feature/m5-feedback-zhuangzijie')]):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('本周完成工作', level=2)
for h, items in [
    ('1. 代码同步', ['拉取 origin/dev 最新代码，整合 V3 阶段更新。',
        '解决 FeedbackService、ApiController、App.vue 多处合并冲突。']),
    ('2. 审核日志', ['新建 V10__audit_log.sql 迁移，audit_log 表。',
        '新建 AuditLog.java 实体和 AuditLogRepository.java。',
        'FeedbackService.updateStatus() 自动记录状态变更日志。',
        '新增 GET /api/admin/feedback/{id}/audit-logs 端点。']),
    ('3. 反馈统计', ['新建 FeedbackStats.java 实体（含 DailyTrend）。',
        'FeedbackService.getStats() 聚合查询，含准确率和 7 天趋势。',
        '新增 GET /api/admin/feedback/stats 端点。',
        'AdminPanel.vue 统计面板：卡片 + 柱状图。']),
    ('4. 纠错样本更新', ['dev 已有 CorrectionSampleService，采纳反馈时自动创建纠错样本。',
        'M5 验证数据流正确，无需额外开发。']),
    ('5. 文档与交付', ['更新 uc_feedback.md 与代码对齐。', '编写第四周周报。'])]:
    doc.add_heading(h, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading('遇到的问题', level=2)
for p in ['dev 与 M5 分支多处合并冲突，均以 dev 为主线逐一解决。',
          '文件锁定导致 checkout 失败，通过 git checkout --force 解决。']:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('下周计划', level=2)
for p in ['整理最终交付文档。', '配合团队进行 V3 整体联调测试。', '准备答辩材料。']:
    doc.add_paragraph(p, style='List Number')

doc.add_heading('备注', level=2)
doc.add_paragraph('M5 模块四周迭代全部完成：数据库设计 -> 反馈持久化 -> 管理闭环 -> 审核日志与统计。')
doc.add_paragraph('反馈采纳后的样本更新由 CorrectionSampleService 统一管理。')

path = r'C:\Users\asus\Desktop\软件实训\Campuslens\docs\zhuangzijie_week4_report.docx'
doc.save(path)
print(f'Saved: {path}')
