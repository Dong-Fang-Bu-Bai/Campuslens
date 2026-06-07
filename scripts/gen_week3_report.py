"""从 Markdown 周报生成 Word 文档。"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 标题
title = doc.add_heading('第三周工作周报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息表格
doc.add_heading('基本信息', level=2)
table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
rows_data = [
    ('项目', '内容'),
    ('姓名', '庄子杰'),
    ('模块', 'M5 用户反馈纠错与检索记录统计'),
    ('日期', '2026-06-01 ~ 2026-06-07'),
    ('分支', 'feature/m5-feedback-zhuangzijie'),
]
for i, (k, v) in enumerate(rows_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# 本周完成工作
doc.add_heading('本周完成工作', level=2)

sections = [
    ('1. 代码同步与合并', [
        '拉取 origin/dev 最新代码，整合第三周 V2 全量更新（认证系统、管理员后台、Flyway 迁移、search_record 持久化）。',
        '解决 FeedbackService、ApiController、schema.sql、App.vue 四处合并冲突，保留 M5 独有的搜索端点反馈和前端 UX 改进。',
    ]),
    ('2. 反馈一致性审查', [
        '逐项核对 correct、wrong、uncertain 三类反馈在前端表单、后端校验和用例规约中的定义，确认一致。',
        '审查 feedback 请求五个字段（searchRecordId、predictedLandmarkId、confirmedLandmarkId、feedbackType、comment）的传递链路。',
        '确认 wrong 类型要求 confirmedLandmarkId 的校验规则在后端执行正确，并补充前端 client-side 拦截。',
    ]),
    ('3. 文档对齐', [
        '修正 uc_feedback.md 表名：user_feedback -> feedback，search_records -> search_record。',
        '修正 BR5 状态值：reviewed / adopted -> accepted / ignored，与后端代码保持一致。',
        '确认业务规则 BR1-BR5 与代码实现完全对应。',
    ]),
    ('4. 关联校验验证', [
        '验证 SearchRecordService.exists() 对 searchRecordId 的存在性校验。',
        '验证 SearchRecordService.containsResult() 对 predictedLandmarkId 必须来自 Top-5 快照的校验。',
        '确认 wrong 类型缺少 confirmedLandmarkId 时后端返回 400，前端优先拦截。',
    ]),
    ('5. 管理后台闭环', [
        '确认 GET /api/admin/feedback 返回最近 50 条反馈记录（含 JOIN 查询地标名和用户名）。',
        '确认 POST /api/admin/feedback/{id}/status 支持将状态更新为 pending、accepted 或 ignored，需管理员 token。',
        '确认 AdminPanel.vue 组件展示反馈列表并提供采纳/忽略操作按钮。',
    ]),
    ('6. V2 主流程测试', [
        '修复两个测试（admin 端点补充 token 鉴权、反馈校验测试适配 searchRecordId 存在性断言）。',
        '新增 feedbackWrongRejectsWithoutConfirmedForValidSearch 测试，覆盖完整场景。',
        '最终 17 项测试全部通过（BUILD SUCCESS）。',
    ]),
]

for heading, items in sections:
    doc.add_heading(heading, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# 遇到的问题
doc.add_heading('遇到的问题', level=2)
problems = [
    '多次 rebase 时旧提交与 dev 新代码冲突，最终采用 merge 策略并以 dev 为主解决。',
    '两个测试因 dev 新增 auth 机制和校验顺序调整而失败，已逐一修复。',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

# 下周计划
doc.add_heading('下周计划', level=2)
plans = [
    '反馈统计数据接口与展示。',
    '审核日志记录。',
    '第四周周报与最终交付文档整理。',
]
for p in plans:
    doc.add_paragraph(p, style='List Number')

# 备注
doc.add_heading('备注', level=2)
doc.add_paragraph('第三周 M5 最小处理闭环已完成：提交反馈 -> 入库 pending -> 管理员审核 accepted/ignored。')
doc.add_paragraph('反馈采纳后的样本更新、统计图表和审核日志写入后续迭代，与第三周 V2 完成记录保持一致。')

output_path = r'C:\Users\asus\Desktop\软件实训\Campuslens\docs\zhuangzijie_week3_report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
