"""从 Markdown 周报生成 Word 文档。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 标题
title = doc.add_heading('第二周工作周报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息表格
doc.add_heading('基本信息', level=2)
table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
rows_data = [
    ('项目', '内容'),
    ('姓名', '庄子杰'),
    ('模块', 'M5 用户反馈纠错与检索记录统计'),
    ('日期', '2026-05-26 ~ 2026-05-31'),
    ('分支', 'feature/m5-feedback-zhuangzijie'),
]
for i, (k, v) in enumerate(rows_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# 本周完成工作
doc.add_heading('本周完成工作', level=2)

sections = [
    ('1. 环境同步', [
        '拉取 origin/dev 最新代码（含 MySQL 支持、算法客户端、深色主题 UI 改版等 18 个新提交）。',
        '将 feature/m5-feedback-zhuangzijie rebase 到最新 dev，解决分叉后强制推送同步远程。',
    ]),
    ('2. 反馈数据持久化', [
        '新建 FeedbackRecord.java 实体 record，映射 feedback 表全部字段（id、searchRecordId、predictedLandmarkId、confirmedLandmarkId、feedbackType、comment、status、createdAt）。',
        '新建 FeedbackRepository.java，使用 JdbcTemplate 实现 CRUD，参照 LandmarkService 的编码风格。支持 save()、findById()、findBySearchRecordId()、findAll()。',
        '改造 FeedbackService.java：注入 FeedbackRepository，submit() 将反馈持久化到 feedback 表，返回真实数据库 ID。保留全部原有校验逻辑（feedbackType 白名单、predictedLandmarkId 和 confirmedLandmarkId 必填校验）。',
    ]),
    ('3. 新增反馈查询接口', [
        'GET /api/admin/feedback：返回全部反馈记录列表，按时间倒序。',
        'GET /api/search/{searchRecordId}/feedback：返回指定检索记录关联的全部反馈。',
    ]),
    ('4. 数据库调整', [
        '修改 database/schema.sql，暂时移除 feedback 表对 search_record 的外键约束。原因：M1 模块暂未持久化检索记录，外键会导致反馈入库失败。待 M1 接入后恢复。',
    ]),
    ('5. 前端反馈表单优化', [
        '新增 feedbackSubmitting 提交中状态和 feedbackError 独立错误提示。',
        '提交按钮增加 loading 动画和 disabled 状态。',
        '提交成功后自动重置表单（清空补充说明、恢复反馈类型）。',
        '错误与成功消息分离显示，各自使用独立样式。',
    ]),
    ('6. 测试补充', [
        '在 ApiControllerTest.java 新增 6 个反馈相关测试：空 feedbackType 返回 400、无效 feedbackType 返回 400、wrong 类型缺少 confirmedLandmarkId 返回 400、uncertain 类型成功提交、管理端反馈列表返回数组、指定检索记录反馈列表返回数组。',
    ]),
]

for heading, items in sections:
    doc.add_heading(heading, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# 遇到的问题
doc.add_heading('遇到的问题', level=2)
problems = [
    'M1 尚未实现 search_record 表持久化，feedback 表外键会导致插入失败。临时方案：移除 FK 约束，待 M1 接入后恢复。',
    '代码 rebase 后分支分叉，需 force-with-lease 推送。',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

# 下周计划
doc.add_heading('下周计划', level=2)
plans = [
    '与 M1 协调 search_record 入库进度，恢复 FK 约束。',
    '实现反馈状态流转（pending -> reviewed -> adopted）。',
    '实现 GET /api/admin/feedback 的管理端页面雏形。',
    '对接 M4 确认反馈入口交互细节。',
]
for p in plans:
    doc.add_paragraph(p, style='List Number')

# 备注
doc.add_heading('备注', level=2)
doc.add_paragraph('当前版本反馈已写入 MySQL feedback 表，检索记录暂由 M1 的 AtomicLong 生成。')
doc.add_paragraph('第二周核心目标 "searchRecordId 与反馈数据的衔接" 已达成。')

output_path = r'C:\Users\asus\Desktop\软件实训\Campuslens\docs\zhuangzijie_week2_report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
